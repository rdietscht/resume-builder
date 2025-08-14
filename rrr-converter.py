import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Section headers to look for when scanning the document.
HEADER_OPTS = [
    "OBJECTIVE",
    "EDUCATION",
    "TECHNICAL SKILLS",
    "RELEVANT EXPERIENCE",
    "PROJECTS/RESEARCH",
    "ADDITIONAL HIGHLIGHTS",
    "AWARDS/HONORS"
]

# Currently supported resume content types.
CONTENT_TYPES = [
    "BULLETED",
    "DESCRIPTION",
    "NEWLINE"
]

class RRR_Parameters:

    def __init__(self):

        # TODO - Should add the ability to be able to specify
        # the user information directly in the parsed document
        # using some special syntax.

        # REMOVED FEATURE
        # for header in HEADER_OPTS:
        #     self.included_headers.append(header)

        # Personal User information.
        self.user_name = '[Your Name]'
        self.location = 'Richmond, VA'
        self.email = 'email@emailservice.com'
        self.phone = '(XXX) XXX-XXXX'

        # Initialize parse location.
        self.parse_path = "\\Users\\rdiet\\Documents\\Old_Resume.docx"

        # Initialize save location.
        self.save_path = "\\Users\\rdiet\\Documents\\New_Resume.docx"

    def set_parse_path(self, new_path):

        # Replace linux path style if needed.
        if ("/" in new_path):
            new_path = new_path.replace("/", "\\")

        self.parse_path = new_path
        print("Saved new parse path successfully")

    def set_save_path(self, new_path):

        # Replace linux path style if needed.
        if ("/" in new_path):
            new_path = new_path.replace("/", "\\")

        self.save_path = new_path
        print("Saved new save path successfully")

    def remove_header(self, header):

        # Ensure the user's header is a valid option and is not already removed from self.included_headers.
        if (not header in HEADER_OPTS):
            print(f"Sorry, {header} is not a valid option.")
        elif (not header in self.included_headers):
            print(f"Header {header} already removed from the header options.")
        else:
            self.included_headers.remove(header)
            print(f"Successfully removed {header} from the included header list.")

    def add_header(self, header):

        if (not header in HEADER_OPTS):
            print(f"Sorry, {header} is not a valid option.")
        elif (header in self.included_headers):
            print(f"Header {header} is already included within the header options list.")
        else:
            self.included_headers.append(header)
            print(f"Successfully added {header} to included headers list.")

    def show_included_headers(self):

        # Format included header options.
        formatted_str = "  Included Headers: (+/included, -/removed)\n"
        for header in HEADER_OPTS:
            if (header in HEADER_OPTS):
                formatted_str += '    + ' if header in self.included_headers else '    - '
                formatted_str += f"{header.capitalize()}\n"
        formatted_str += "\n"
        return formatted_str

    def show_parse_path(self):
        return f"  Resume Content Path: {self.parse_path:>47.47}\n"

    def show_save_path(self):
        return f"  Document Save Location: {self.save_path:>44.44}\n"

    def show_personal_info(self):
        return (
            f"  Personal User Information:\n\
    1. (NAME) {self.user_name}\n\
    2. (LOCATION) {self.location}\n\
    3. (EMAIL) {self.email}\n\
    4. (PHONE) {self.phone}\n"
        )

    def __str__(self):
        formatted_str = f"{'CURRENT PARAMETERS':<25}\n"

        # formatted_str += self.show_included_headers() # REMOVED FEATURE

        formatted_str += self.show_personal_info() + "\n"
        formatted_str += self.show_parse_path() + "\n"
        formatted_str += self.show_save_path() + "\n"

        return formatted_str

# Stores all the content and formatting info for each section of a single resume document.
class File_Handle:

    def __init__(self):
        self.sections = [] # sections stores Section instances

    def __str__(self): # debugging only

        formatted_str = "  File Sections:\n"
        for section in self.sections:
            formatted_str += f"\t{str(section)}\n"

        return formatted_str


# Stores any subsections and their content.
class File_Section:

    def __init__(self, title):
        self.title = title
        self.content_list = [] # stores array of File_Sub | File_Content instances

    def __str__(self): # debugging only

        formatted_str = f"{self.title}:\n"
        for content in self.content_list:
            formatted_str += f"\t  {str(content)}\n"

        return formatted_str


class File_Sub:

    def __init__(self, title):
        self.title = title

    def __str__(self):

        return f"{self.title}"


# Stores a single unit of file content, either a bulleted list or a paragraph description.
class File_Content:

    def __init__(self, content_type, content):
        self.type = content_type
        self.content = content # of type: string[] (BULLETED) | string (DESCRIPTION/NEWLINE)

    def __str__(self):

        return f"(TYPE: {self.type}, CONTENT: \"{str(self.content):.27}...\")"


# UTILITY FUNCTIONS
def define_parameters():

    # Greet the user.
    wipe_terminal()
    print(f"{'Welcome':^25}")
    print(f"{'RAD-Resume-Retrofitter':^25}")
    print()

    # Ask to configure settings for an RRR_Parameters instance.
    params = RRR_Parameters()
    user_input = ""
    while (user_input != "N" and user_input != "NEXT"):

        # EARLY EXIT - The user has quit the program.
        if (user_input == "Q" or user_input == "QUIT"):
            print("Goodbye!")
            quit()

        # Handle modification of included headers list. # REMOVED FEATURE
        # if (user_input == "H" or user_input == "HEADER"):
        #     wipe_terminal()
        #     while(user_input != 'Q' and user_input != 'QUIT'):

        #         # ADDING A HEADER.
        #         if (user_input == "A" or user_input == "ADD"):

        #             wipe_terminal()
        #             print(params.show_included_headers())
        #             print("Which header?")
        #             print()
        #             print("Type name of header: ", end="")
        #             user_input = input().upper()
        #             wipe_terminal()
        #             params.add_header(user_input)

        #             # Once a header is successfully added, return to the main menu.
        #             break

        #         # REMOVING A HEADER.
        #         if (user_input == "R" or user_input == "REMOVE"):

        #             wipe_terminal()
        #             print(params.show_included_headers())
        #             print("Which header?")
        #             print()
        #             print("Type name of header: ", end="")
        #             user_input = input().upper()
        #             wipe_terminal()
        #             params.remove_header(user_input)

        #             # Once a header is successfully added, return to the main menu.
        #             break

        #         print("Would you like to add/remove a header?")
        #         print("(Q/Quit to go back to previous menu, A/Add to add, R/Remove to remove..)")
        #         print()
        #         print("Enter: ", end="")

        #         user_input = input().upper()

        #     print()

        # Handle changing of personal user info.
        if (user_input == "U" or user_input == "USER"):

            wipe_terminal()

            # Prompt the user to select the number corresponding to the user info being modified.
            while (user_input != "Q" and user_input != "QUIT"):

                if (user_input == "1"): # NAME
                    wipe_terminal()
                    print(f"Current Value: {params.user_name}")
                    print()
                    print("New Value: ", end="")
                    params.user_name = input()
                    wipe_terminal()
                    print(f"Value \"{params.user_name}\" saved successfully")
                    print()
                    break
                elif (user_input == "2"): # LOCATION
                    wipe_terminal()
                    print(f"Current Value: {params.location}")
                    print()
                    print("New Value: ", end="")
                    params.location = input()
                    wipe_terminal()
                    print(f"Value \"{params.location}\" saved successfully")
                    print()
                    break
                elif (user_input == "3"): # EMAIL
                    wipe_terminal()
                    print(f"Current Value: {params.email}")
                    print()
                    print("New Value: ", end="")
                    params.email = input()
                    wipe_terminal()
                    print(f"Value \"{params.email}\" saved successfully")
                    print()
                    break
                elif (user_input == "4"): # PHONE
                    wipe_terminal()
                    print(f"Current Value: {params.phone}")
                    print()
                    print("New Value: ", end="")
                    params.phone = input()
                    wipe_terminal()
                    print(f"Value \"{params.phone}\" saved successfully")
                    print()
                    break

                # Show users the currently saved data.
                print(params.show_personal_info())
                print()
                print("Please enter the number corresponding to the information type being modified.")
                print("(Q/Quit to cancel and go back to the previous menu..)")
                print("Enter Number: ", end="")
                user_input = input().upper()

        # Handle changing of parse/save location.
        if (user_input == "P" or user_input == "PARSE" or user_input == "S" or user_input == "SAVE"):

            wipe_terminal()

            parse_flag = False # indicates the user is saving a new parse path

            # Save to appropriate place.
            if (user_input == "P" or user_input == "PARSE"):
                parse_flag = True

            print("Please type in the full path for the file.")
            print("(Q/Quit to cancel and go back to previous menu..)")
            print()
            print("New Path: ", end="")
            user_input = input()

            while (user_input.upper() != 'Q' and user_input.upper() != 'QUIT'):

                # Only accept non-empty path settings.
                if (user_input != ""):
                    if (parse_flag):
                        wipe_terminal()
                        params.set_parse_path(user_input)
                        break
                    else:
                        wipe_terminal()
                        params.set_save_path(user_input)
                        break

                print("Please type in the full path for the file.")
                print("(Q/Quit to cancel and go back to previous menu..)")
                print()
                print("New Path: ", end="")
                user_input = input()

            print()

        # Show current status of the parameters.
        print(params)

        print("Please Select an option:")
        print("(U/User to modify personal information, P/Parse to modify parse location, S/Save to modify save location..)")
        print("(Q/Quit to cancel, N/Next to continue with the saved parameters..)")
        print()
        print("Enter: ", end="")
        user_input = input().upper()

    return params

 
def scan_formatted_document(parse_path):

    print("Scanning document content...")
    print()

    # TODO - Performance can def be improved by interpreting each line/character as it is read.
    r_lines = [] # stores a list of strings representing each line in the file.
    with open(parse_path, 'r', encoding='utf-8') as fh:

        # Gather each line of the file.
        line = fh.readline()
        while line != "":
            r_lines.append(line)
            line = fh.readline()

        # Show the user content of the file scanned.
        print(f"{'RESUME CONTENT SCANNED':^100}")
        print('=' * 100)
        print()
        for r_line in r_lines:
            if (len(r_line) > 100):
                print(f"{r_line:<97.97}...")
            else:
                print(f"{r_line:<100.100}")
        print()
        print('=' * 100)
        print()

    # Go through each line, adding to the File_Handle sections array.
    file_handle = File_Handle()
    for line_i in range(len(r_lines)):

        # CASE - Section header (#) encountered.
        if ("#" in r_lines[line_i]):

            # Consume the section title.
            entered_header = False
            section_title = ""
            for char in r_lines[line_i]:

                # Ignore the initial header character.
                if (char == "#"):
                    if (not entered_header): # we have entered the header.
                        entered_header = True
                    else:
                        section_title += char
                elif (entered_header):

                    if (char == '\n'):
                        break

                    section_title += char

            # Consume the section's content, including any sub-sections, bulleted points, and paragraphs.
            fs = File_Section(section_title)
            print(f"Parsing section starting at line {line_i} (\"{section_title}\")")
            line_i = consume_section(line_i + 1, r_lines, fs)
            print(f"Consumed section; resuming search for sections from {line_i}")
            file_handle.sections.append(fs)

    # Return the file_handle with all its stored content.
    return file_handle


def consume_section(index, lines, fs):

    # Keep interpreting lines until the end of the section header is encountered or the lines end.
    while index < len(lines) and lines[index] != '\n':

        # A section content may consist of sub-section headers or content (either descriptions or bulleted lists)
        # In other words, we have the following 3 entrance states:
        #   $    - Create a new sub-section with the provided title. End any previous sub-section we were on.
        #   """  - Create and append a "DESCRIPTION" File_Content instance.
        #   ``   - Create and append a "BULLETED" File_Content instance.
        i = 0 # character index
        while i < len(lines[index]):

            if (i + 2 < len(lines[index]) and lines[index][i:i + 3] == '"""'): # DESCRIPTION ENTRANCE
                # print("Encountered description!")

                i += 3 # skip past the three quotes

                # Keep consuming regular text characters until the ending (""") token is encountered.
                d_content = ""
                closed = False
                while i + 2 < len(lines[index]): # TODO - FOR THIS AND DESC, THIS CONSTRAINT KEEPS MULTI-LINE SYNTAX FROM WORKING. MODIFY THIS SO THAT END DELIMITERS MAY WORK ON A SEPARATE LINE.

                    # Terminate loop when ending description token is found. Move index past the characters.
                    if (lines[index][i:i + 3] == '"""'):
                        closed = True
                        i += 3
                        break

                    # Consume the character and move to the next.
                    d_content += lines[index][i]
                    i += 1

                # EARLY EXIT - The user template left an open description before ending the line.
                if (not closed):
                    print()
                    print(f"ERR: Found unclosed description content on line {index}")
                    print()
                    quit()

                fs.content_list.append(File_Content(CONTENT_TYPES[1], d_content))

            elif (lines[index][i] == '`'): # BULLETED ENTRANCE
                # print("Encountered bulleted list!")

                i += 1 # skip past token

                # Keep consuming text characters, including delimiters.
                b_raw_content = ""
                closed = False
                while (i < len(lines[index])): # TODO - FOR THIS AND DESC, THIS CONSTRAINT KEEPS MULTI-LINE SYNTAX FROM WORKING. MODIFY THIS SO THAT END DELIMITERS MAY WORK ON A SEPARATE LINE.

                    # Terminate loop when ending bulleted token is found. Move index past end token.
                    if (lines[index][i] == "`"):
                        closed = True
                        i += 1
                        break

                    # Consume the character and move to the next.
                    b_raw_content += lines[index][i]
                    i += 1

                # EARLY EXIT - The user template left an open description before ending the line.
                if (not closed):
                    print()
                    print(f"ERR: Found unclosed list content on line {index}")
                    print()
                    quit()

                # After raw content is consumed, create an array of string values using the delimiters.
                b_content = b_raw_content.split('*')[1:-1]
                fs.content_list.append(File_Content(CONTENT_TYPES[0], b_content))

            elif (lines[index][i] == '$'): # NAMED SUB-HEADER ENTRANCE
                # print("Encountered sub-section!") # DEBUGGING!

                i += 1 # skip past token

                # Consume the characters of the sub-section title.
                sub_title = ""
                closed = False
                while (i < len(lines[index])):

                    # Encountered sub-header end. Move index past the end token.
                    if (lines[index][i] == '$'):
                        closed = True
                        i += 1
                        break

                    # Add to the sub title.
                    sub_title += lines[index][i]
                    i += 1

                # EARLY EXIT - The user template left an open description before ending the line.
                if (not closed):
                    print()
                    print(f"ERR: Found unclosed subheader on line {index}")
                    print()
                    quit()

                fs.content_list.append(File_Sub(sub_title))

            elif (lines[index][i] == '\n'):

                nl_content = File_Content(CONTENT_TYPES[2], '_LINE_BREAK_')
                fs.content_list.append(nl_content)
                i += 1

            else:
                print()
                print(f"ERR: Encountered unexpected token ( {lines[index][i]} ) at position ({index}, {i})")
                print()
                quit()

        # Progress to next line.
        index += 1

    return index


"""Create a .docx file using the file content and parameters.

@params
    content - A File_Content instance with the contents to be included in the formatted docx.
    params - The custom user parameters to include when formatting the document.
"""
def create_formatted_document(content: File_Handle, params: RRR_Parameters):

    print("Creating + Formatting document...")

    # Create the Word document.
    doc = Document()

    # Add personal info at the top of the document.
    pp = doc.add_paragraph() # hehe
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp_format = pp.paragraph_format
    pp_format.space_before = Pt(0)
    pp_format.space_after = Pt(0)
    pp_format.line_spacing = 1.0
    name_run = pp.add_run()
    font = name_run.font
    font.bold = True
    font.name = 'Calibri'
    font.size = Pt(12) 
    name_run.text = params.user_name + "\n"

    lep_run = pp.add_run()
    font = lep_run.font
    font.name = 'Calibri'
    font.size = Pt(7)
    lep_run.text = params.location + '\n' + params.email + '\n' + params.phone

    # Write each section with their associated content.
    for section in content.sections:

        section: File_Section = section

        # # Skip any sections omitted by the user parameters. # REMOVED OLD CONSTRAINT
        # if (not section.title.upper() in params.included_headers):
        #     print(f"Skipping section write: \"{section.title}\"...")
        #     continue

        # FORMATTING FOR SECTIONS.
        header = doc.add_paragraph()
        p_format = header.paragraph_format
        p_format.space_before = Pt(1)
        p_format.space_after = Pt(2)
        p_format.left_indent = Pt(2)
        run = header.add_run()
        run.bold = True
        font = run.font
        font.color.rgb = RGBColor(0,0,0)
        font.name = 'Calibri'
        font.size = Pt(12)
        run.text = section.title

        # Adding a bottom border (GPT-CODE: MODIFY CAREFULLY)
        header_xml = header._p
        pPr = header_xml.get_or_add_pPr()

        # Create a new border element
        border = OxmlElement('w:pBdr')

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')  # Border style
        bottom.set(qn('w:sz'), '4')         # Border size
        bottom.set(qn('w:space'), '0')      # Space
        bottom.set(qn('w:color'), '000000')  # Border color (black)

        # Add the border to the paragraph
        border.append(bottom)
        pPr.append(border)

        # Write each of the sections content.
        paragraph = None
        run = None
        for content in section.content_list:

            # Switch between each content type and sub-headers.
            content_type = type(content)
            if (content_type == File_Sub): # SUB-HEADERS
                if (paragraph == None):
                    paragraph = doc.add_paragraph()
                    p_format = paragraph.paragraph_format
                    p_format.line_spacing = 1.0
                    p_format.space_before = Pt(0)
                    p_format.space_after = Pt(0)
                    run = paragraph.add_run()
                else:
                    run = paragraph.add_run()
                font = run.font
                font.color.rgb = RGBColor(0,0,0)
                font.name = 'Calibri'
                font.size = Pt(11)
                font.bold = True
                run.text += content.title
            elif (content_type == File_Content and content.type == CONTENT_TYPES[1]): # DESCRIPTIONS
                if (paragraph == None):
                    paragraph = doc.add_paragraph()
                    p_format = paragraph.paragraph_format
                    p_format.line_spacing = 1.0
                    p_format.space_before = Pt(0)
                    p_format.space_after = Pt(0)
                    run = paragraph.add_run()
                else:
                    run = paragraph.add_run()
                font = run.font
                font.color.rgb = RGBColor(0,0,0)
                font.name = 'Calibri'
                font.bold = False # TODO - FIX THIS! ALL RUNS SHARE THE SAME SETTINGS!
                font.size = Pt(11)
                run.text += content.content
            elif (content_type == File_Content and content.type == CONTENT_TYPES[0]): # BULLETED
                for bullet_content in content.content:
                    bullet = doc.add_paragraph(style='List Bullet')
                    b_format = bullet.paragraph_format
                    b_format.line_spacing = 1.0
                    b_format.space_before = Pt(0)
                    b_format.space_after = Pt(0)
                    run = bullet.add_run()
                    run.text = bullet_content
                paragraph = None # reset the paragraph after bullet list items
            elif (content_type == File_Content and content.type == CONTENT_TYPES[2]): # NEWLINE
                if (paragraph != None):
                    # run = paragraph.add_run()
                    # run.add_break()
                    run.text += "\n"
            else:
                print()
                print(f"ERR: An invalid type was found when writing section content to the document ({content_type})")
                print()
                quit()


    # Print a confirmation to show the operation was successful.
    doc.save(params.save_path)
    print(f"Done - Word document saved in: {params.save_path}")


def wipe_terminal():
    # wipe terminal (either Windows/Linux)
    os.system('cls' if os.name == 'nt' else 'clear')

# Main script
if __name__ == '__main__':

    # Ask the user to define parameters (i.e., path to
    # formatted document, path to created document,
    # headers to include w/ options, etc...).
    parameters = define_parameters()

    # Obtain file information and content details. Log any errors encountered.
    f_handle = scan_formatted_document(parameters.parse_path)
    print(f_handle) # DEBUGGING!

    # Create a new word file using configured settings.
    create_formatted_document(f_handle, parameters)
